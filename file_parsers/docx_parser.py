"""
DOCX parser — extracts text from Microsoft Word documents using python-docx.
"""

import io
import logging
from .base import BaseParser

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Extract text from .docx files using python-docx."""

    supported_mime_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ]

    def parse(self, data: bytes, filename: str = "") -> dict:
        try:
            from docx import Document  # python-docx
        except ImportError:
            raise RuntimeError(
                "python-docx is not installed. Run: pip install python-docx"
            )

        try:
            doc = Document(io.BytesIO(data))
        except Exception as exc:
            logger.error(f"DOCX parsing failed for '{filename}': {exc}")
            raise ValueError(f"Failed to parse DOCX: {exc}") from exc

        parts: list[str] = []

        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Extract table cell text
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(c for c in row_cells if c)
                if row_text:
                    parts.append(row_text)

        full_text = "\n".join(parts)
        return {
            "text": full_text,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }
